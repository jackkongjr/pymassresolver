#!/usr/bin/env python3
import sys
import getopt
import dns.resolver
import re
import docx



savefile = None

def clean(hostname):
    hostname = hostname.strip().replace("\n","")
    #removes ansi color codes
    ansi_escape = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')
    hostname = ansi_escape.sub('', hostname)
    return hostname

def is_valid_hostname(hostname):
    regex = "^((?!-)[A-Za-z0-9-]{1,63}(?<!-)\\.)+[A-Za-z]{2,6}"     
    p = re.compile(regex)
    if (hostname == None):
        return False
    if(re.search(p, hostname)):
        return True
    else:
        return False

def usage():
    u = """
 ____          __  __                 ____                 _                
|  _ \ _   _  |  \/  | __ _ ___ ___  |  _ \ ___  ___  ___ | |_   _____ _ __ 
| |_) | | | | | |\/| |/ _` / __/ __| | |_) / _ \/ __|/ _ \| \ \ / / _ \ '__|
|  __/| |_| | | |  | | (_| \__ \__ \ |  _ <  __/\__ \ (_) | |\ V /  __/ |   
|_|    \__, | |_|  |_|\__,_|___/___/ |_| \_\___||___/\___/|_| \_/ \___|_|   
       |___/                                    
       

Simple script to mass resolve multiple fqdn from standard input or from file

Usage: pymassres.py -f <input_file>


Available commands:
    -h                this message
    -i                reads from stdin   
    -f / --file       file path to read
    -s / --save       save as docx table

It can detect malformed domain and automatically skip it
    

Example, reading directly from sublist3r output:

    sublist3r -d domain.com  |  python3 pymassres.py  -i

or redirecting output from terminal

    cat hosts_file | python pymassres.py -i


"""
    print (u)


def resolve(line):
    c_line = clean(line)
    if is_valid_hostname(c_line):
        try:
            result = dns.resolver.resolve(c_line, 'A')
            ips = []
            for val in result:
                ips.append (val.to_text())
            print (c_line + " " + " ".join(ips) )
            return c_line,ips

        except dns.exception.DNSException:
            return "",""
    else:
        return "",""


def main(argv):
    global savefile
    inputfile = ''
    try:
        opts, args = getopt.getopt(argv,"hif:s:",["file=","save="])
    except getopt.GetoptError:
        usage()
        sys.exit(2)

    if len(opts)==0:
        usage()
        exit(0)

    doc = docx.Document()
    doc.add_heading('Table', 0)
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0].cells
    row[0].text = 'Hostname'
    row[1].text = 'IP'
    done_op = False
    for opt, arg in opts:
        if opt == '-h':
            usage()
            sys.exit()
        elif opt in ("-i"):
            for line in sys.stdin:               
                c,i = resolve(line)
                if c and i :
                    row = table.add_row().cells
                    row[0].text = c
                    row[1].text = " ".join(i)
                    done_op = True

        elif opt in ("-s","--save"):
            savefile = arg

        elif opt in ("-f", "--file"):
            content = []
            inputfile = arg
            try:
                f = open(inputfile,'r')
                content = f.readlines()
            except:
                print("Error accessing file")
                sys.exit(0)
            for line in content:
                c,i = resolve(line)
                if c and i:
                    row = table.add_row().cells
                    row[0].text = c
                    row[1].text = " ".join(i)
                    done_op = True
    if done_op and savefile:
        doc.save(savefile)


if __name__=='__main__':
    try:
        main(sys.argv[1:])
    except KeyboardInterrupt:
        sys.exit(0)